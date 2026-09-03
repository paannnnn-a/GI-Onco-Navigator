from pypdf import PdfWriter

from backend.app.knowledge import (
    ExtractedPage,
    audit_pdf,
    chunk_pages,
    chunk_transcript,
    extract_docx_paragraphs,
    extract_pdf_pages,
    extract_transcript_cues,
    looks_corrupted,
)


def test_corrupt_text_detection() -> None:
    assert looks_corrupted("")
    assert looks_corrupted("δ൭Ҩ" * 30)
    assert looks_corrupted("໇ᆀ൭ҨழႶૌ౦" * 20)
    assert not looks_corrupted("这是一段用于患者教育的正常中文文本。" * 10)
    assert not looks_corrupted("HER2、PD-L1、MSI-H/dMMR 与 β-catenin 是常见医学写法。" * 5)


def test_chunk_preserves_page_locator() -> None:
    pages = [
        ExtractedPage(3, "第一段。" * 30, "pdf_text", False),
        ExtractedPage(4, "第二段。" * 30, "pdf_text", False),
    ]
    chunks = chunk_pages(pages, target_chars=1000)
    assert len(chunks) == 1
    assert chunks[0].page_start == 3
    assert chunks[0].page_end == 4


def test_extract_docx_paragraphs(tmp_path) -> None:
    from docx import Document

    path = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("第一段可用资料")
    document.add_paragraph("")
    document.add_paragraph("第二段可用资料")
    document.save(path)
    blocks = extract_docx_paragraphs(path)
    assert [block.page_number for block in blocks] == [1, 3]
    assert all(block.method == "docx_paragraph" for block in blocks)


def test_pdf_audit_contains_metrics_but_no_document_text(tmp_path) -> None:
    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as output:
        writer.write(output)

    result = audit_pdf(path)
    assert result["pages"] == 1
    assert result["pages_needing_ocr"] == [1]
    assert result["readable_text_pages"] == 0
    assert "text" not in result


def test_pdf_extraction_uses_batch_ocr_when_available(tmp_path) -> None:
    path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as output:
        writer.write(output)

    class BatchOcr:
        def __init__(self) -> None:
            self.requested: list[int] = []

        def recognize_pdf_pages(self, _path, page_numbers):
            self.requested = list(page_numbers)
            return {number: "本地识别出的患者教育文本。" * 5 for number in self.requested}

        def recognize_pdf_page(self, _path, _page_number):
            raise AssertionError("single-page OCR should not be used")

    engine = BatchOcr()
    pages = extract_pdf_pages(path, engine)
    assert engine.requested == [1, 2]
    assert all(page.method == "ocr" and not page.needs_ocr for page in pages)


def test_transcript_preserves_video_timestamp(tmp_path) -> None:
    path = tmp_path / "expert.srt"
    path.write_text(
        "1\n00:00:32,100 --> 00:00:38,900\n复诊前请整理已有资料。\n\n"
        "2\n00:00:39,000 --> 00:00:45,000\n具体情况需要与诊疗团队确认。\n",
        encoding="utf-8",
    )
    cues = extract_transcript_cues(path)
    chunks = chunk_transcript(cues)
    assert len(cues) == 2
    assert chunks[0].start_seconds == 32
    assert chunks[0].end_seconds == 45
    assert "诊疗团队" in chunks[0].text
